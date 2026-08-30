#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""End-to-end test: synthesise sources, run every stage, validate the output.

    python3 tools/contract-corpus/test_pipeline.py

Builds a small .docx — with real Word auto-numbering in numbering.xml, a custom
`CenterHeading` style, a table, and a header and footer — plus a small fake
Azure `prebuilt-layout` JSON, runs ingest → clean → segment → annotate → build →
qc over them, and checks the corpus against validate_corpus.py.

No network, no fixtures on disk, nothing written outside a temporary directory.
"""
import gzip
import json
import os
import shutil
import sys
import tempfile
import unittest

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)

import annotate                                     # noqa: E402
import build as pipeline                            # noqa: E402
import validate_corpus                              # noqa: E402

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

GOVERNING_LAW = ("สัญญาฉบับนี้อยู่ภายใต้บังคับและตีความตามกฎหมายแห่งราชอาณาจักรไทย "
                 "และให้ศาลไทยเป็นผู้มีอำนาจพิจารณาข้อพิพาท")


# ── A .docx with real auto-numbering ─────────────────────────────────────────

def _oxml(tag):
    from docx.oxml import OxmlElement
    return OxmlElement(tag)


def _add_abstract(numbering, abstract_id, levels):
    """levels: [(numFmt, lvlText, start)] — one w:lvl each."""
    element = _oxml("w:abstractNum")
    element.set(W + "abstractNumId", str(abstract_id))
    for ilvl, (fmt, text, start) in enumerate(levels):
        lvl = _oxml("w:lvl")
        lvl.set(W + "ilvl", str(ilvl))
        for tag, value in (("w:start", str(start)), ("w:numFmt", fmt),
                           ("w:lvlText", text)):
            child = _oxml(tag)
            child.set(W + "val", value)
            lvl.append(child)
        element.append(lvl)
    first_num = numbering.find(W + "num")
    if first_num is not None:
        first_num.addprevious(element)
    else:
        numbering.append(element)


def _add_num(numbering, num_id, abstract_id):
    element = _oxml("w:num")
    element.set(W + "numId", str(num_id))
    ref = _oxml("w:abstractNumId")
    ref.set(W + "val", str(abstract_id))
    element.append(ref)
    numbering.append(element)


def _number(paragraph, num_id, ilvl):
    """Put the paragraph under Word auto-numbering — the number stays OUT of
    the text, exactly as in the real contracts (DESIGN.md §3.3)."""
    properties = paragraph._p.get_or_add_pPr()
    num_pr = _oxml("w:numPr")
    for tag, value in (("w:ilvl", str(ilvl)), ("w:numId", str(num_id))):
        child = _oxml(tag)
        child.set(W + "val", value)
        num_pr.append(child)
    properties.append(num_pr)
    return paragraph


def make_docx(path):
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_BREAK

    document = Document()
    numbering = document.part.numbering_part.element
    _add_abstract(numbering, 900, [("decimal", "%1.", 1)])                  # parties
    _add_abstract(numbering, 901, [("thaiLetters", "%1.", 1)])              # recitals
    _add_abstract(numbering, 902, [("decimal", "%1.", 1),                   # operative
                                   ("decimal", "%1.%2.", 1),
                                   ("decimal", "(%1)", 1)])
    for num_id, abstract_id in ((900, 900), (901, 901), (902, 902)):
        _add_num(numbering, num_id, abstract_id)

    styles = document.styles
    if "CenterHeading" not in [s.name for s in styles]:
        heading = styles.add_style("CenterHeading", WD_STYLE_TYPE.PARAGRAPH)
        heading.base_style = styles["Normal"]

    section = document.sections[0]
    section.header.paragraphs[0].text = "สัญญาจ้างทำของ — ฉบับลงนาม"
    section.footer.paragraphs[0].text = "หน้า 1 จาก 2"

    document.add_paragraph("สัญญาจ้างทำของ", style="CenterHeading")
    document.add_paragraph(
        "สัญญาฉบับนี้ทำขึ้น ณ กรุงเทพมหานคร เมื่อวันที่ ๑๑ มีนาคม ๒๕๖๗")
    _number(document.add_paragraph(
        "บริษัท ก จำกัด สำนักงานตั้งอยู่เลขที่ ๑ ถนนสาทร ซึ่งต่อไปนี้เรียกว่า “ผู้ว่าจ้าง” ฝ่ายหนึ่ง"),
        900, 0)
    _number(document.add_paragraph(
        "บริษัท ข จำกัด สำนักงานตั้งอยู่เลขที่ ๒ ถนนสีลม ซึ่งต่อไปนี้เรียกว่า “ผู้รับจ้าง” อีกฝ่ายหนึ่ง"),
        900, 0)
    _number(document.add_paragraph(
        "โดยที่ผู้ว่าจ้างประสงค์จะว่าจ้างให้ผู้รับจ้างดำเนินการพัฒนาระบบสารสนเทศ"), 901, 0)
    _number(document.add_paragraph(
        "โดยที่ผู้รับจ้างมีความรู้ความชำนาญและตกลงรับจ้างตามเงื่อนไขในสัญญานี้"), 901, 0)

    _number(document.add_paragraph("คำนิยาม", style="Heading 1"), 902, 0)
    _number(document.add_paragraph(
        "“งาน” หมายความว่า งานพัฒนาระบบสารสนเทศตามขอบเขตของงานในเอกสารแนบท้าย"), 902, 1)

    _number(document.add_paragraph("ค่าจ้างและการชำระเงิน", style="Heading 1"), 902, 0)
    _number(document.add_paragraph(
        "ผู้ว่าจ้างตกลงชำระค่าจ้างเป็นงวดตามตารางด้านล่างนี้ ภายในสามสิบวันนับแต่วันที่ได้รับใบแจ้งหนี้"),
        902, 1)
    table = document.add_table(rows=2, cols=3)
    for row, cells in enumerate((("งวดที่", "งานที่ส่งมอบ", "ค่าจ้าง"),
                                 ("๑", "ส่งมอบระบบและเอกสารคู่มือ", "๑๐๐,๐๐๐ บาท"))):
        for column, text in enumerate(cells):
            table.rows[row].cells[column].text = text

    _number(document.add_paragraph("การบอกเลิกสัญญา", style="Heading 1"), 902, 0)
    _number(document.add_paragraph(
        "ในกรณีที่ผู้รับจ้างผิดนัดไม่ส่งมอบงานภายในกำหนด ผู้ว่าจ้างมีสิทธิบอกเลิกสัญญาได้ทันที "
        "โดยไม่ต้องบอกกล่าวล่วงหน้า"), 902, 1)
    document.add_paragraph(
        "การบอกเลิกสัญญาตามวรรคก่อนไม่กระทบต่อสิทธิเรียกร้องค่าเสียหายที่เกิดขึ้นแล้ว")
    _number(document.add_paragraph(
        "ผู้ว่าจ้างอาจบอกเลิกสัญญาเมื่อใดก็ได้โดยบอกกล่าวเป็นหนังสือล่วงหน้าไม่น้อยกว่าสามสิบวัน"),
        902, 2)

    _number(document.add_paragraph("กฎหมายที่ใช้บังคับ", style="Heading 1"), 902, 0)
    _number(document.add_paragraph(GOVERNING_LAW), 902, 1)

    break_paragraph = document.add_paragraph()
    break_paragraph.add_run().add_break(WD_BREAK.PAGE)
    document.add_paragraph("ลงชื่อ ..................... ผู้ว่าจ้าง")
    document.add_paragraph("ลงชื่อ ..................... ผู้รับจ้าง")
    document.add_paragraph("เอกสารแนบท้าย ก ขอบเขตของงาน")
    document.add_paragraph("   ")                   # empty after cleaning
    document.save(path)


# ── A fake Azure prebuilt-layout JSON ────────────────────────────────────────

def make_azure(path):
    """Build the JSON the way Azure lays it out: one `content` string, every
    paragraph and table cell carrying a span into it, words with confidences."""
    items = [
        ("สัญญาไม่เปิดเผยข้อมูล / Non-Disclosure Agreement", "title", 1, 0.99),
        ("สัญญาจ้างทำของ — ฉบับลงนาม", "pageHeader", 1, 0.99),
        ("สัญญาฉบับนี้ทำขึ้นที่กรุงเทพมหานคร เมื่อวันที่ ๒ พฤศจิกายน ๒๕๖๖", None, 1, 0.98),
        ("ระหว่าง บริษัท ค จำกัด ซึ่งต่อไปนี้เรียกว่า “ผู้เปิดเผยข้อมูล” ฝ่ายหนึ่ง", None, 1, 0.97),
        ("ข้อ ๑ คำนิยาม", "sectionHeading", 1, 0.99),
        ("๑.๑ “ข้อมูลอันเป็นความลับ” หมายความว่า ข้อมูลใด ๆ ที่ผู้เปิดเผยข้อมูลส่งมอบให้แก่ผู้รับข้อมูล",
         None, 1, 0.96),
        ("ข้อ ๒ การรักษาความลับ", "sectionHeading", 1, 0.99),
        ("๒.๑ ผู้รับข้อมูลตกลงเก็บรักษาข้อมูลอันเป็นความลับไว้เป็นความลับ และจะไม่เปิดเผยต่อบุคคลภายนอก",
         None, 1, 0.85),
        ("2.2 The Receiving Party shall keep the Confidential Information confidential "
         "and shall not disclose it to any third party.", None, 2, 0.99),
        ("ข้อ ๓ กฎหมายที่ใช้บังคับ", "sectionHeading", 2, 0.99),
        ("๓.๑ " + GOVERNING_LAW, None, 2, 0.98),
        ("ลงชื่อ ..................... ผู้เปิดเผยข้อมูล", None, 2, 0.94),
        ("หน้า 2 จาก 2", "pageFooter", 2, 0.99),
    ]
    table_cells = [("ลำดับ", 0, 0), ("รายการข้อมูล", 0, 1),
                   ("๑", 1, 0), ("แผนธุรกิจและงบการเงิน", 1, 1)]

    content, paragraphs, words, pages = [], [], [], {}
    offset = 0

    def emit(text, page, confidence):
        """Append text to `content`, returning its span, and make words."""
        nonlocal offset
        span = {"offset": offset, "length": len(text)}
        content.append(text)
        cursor = offset
        for word in text.split(" "):
            if word:
                words.append({"content": word,
                              "span": {"offset": cursor, "length": len(word)},
                              "confidence": confidence})
            cursor += len(word) + 1
        pages.setdefault(page, []).append(span)
        offset += len(text) + 1                     # the "\n" between paragraphs
        return span

    for text, role, page, confidence in items[:8]:
        span = emit(text, page, confidence)
        para = {"content": text, "spans": [span],
                "boundingRegions": [{"pageNumber": page, "polygon": [0, 0, 1, 0, 1, 1, 0, 1]}]}
        if role:
            para["role"] = role
        paragraphs.append(para)

    # the table, and its cell paragraphs — Azure lists cells in `paragraphs[]`
    # as well, and the ingester must not emit them twice
    table_start = offset
    cells = []
    for text, row, column in table_cells:
        span = emit(text, 1, 0.99)
        cells.append({"rowIndex": row, "columnIndex": column, "content": text,
                      "spans": [span],
                      "boundingRegions": [{"pageNumber": 1, "polygon": [0, 0, 1, 0, 1, 1, 0, 1]}]})
        paragraphs.append({"content": text, "spans": [span],
                           "boundingRegions": [{"pageNumber": 1, "polygon": []}]})
    table = {"rowCount": 2, "columnCount": 2, "cells": cells,
             "spans": [{"offset": table_start, "length": offset - table_start}],
             "boundingRegions": [{"pageNumber": 1, "polygon": []}]}

    for text, role, page, confidence in items[8:]:
        span = emit(text, page, confidence)
        para = {"content": text, "spans": [span],
                "boundingRegions": [{"pageNumber": page, "polygon": []}]}
        if role:
            para["role"] = role
        paragraphs.append(para)

    page_records = []
    for number, spans in sorted(pages.items()):
        start = min(s["offset"] for s in spans)
        end = max(s["offset"] + s["length"] for s in spans)
        page_records.append({
            "pageNumber": number,
            "spans": [{"offset": start, "length": end - start}],
            "words": [w for w in words
                      if start <= w["span"]["offset"] < end],
        })

    payload = {
        "status": "succeeded",
        "analyzeResult": {
            "apiVersion": "2024-11-30",
            "modelId": "prebuilt-layout",
            "stringIndexType": "utf16CodeUnit",
            "content": "\n".join(content),
            "pages": page_records,
            "paragraphs": paragraphs,
            "tables": [table],
        },
    }
    with open(path, "w", encoding="utf-8") as handle:
        json.dump(payload, handle, ensure_ascii=False)


# ── The test ─────────────────────────────────────────────────────────────────

class PipelineTest(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        cls.tmp = tempfile.mkdtemp(prefix="jtcorpus-test-")
        cls.raw = os.path.join(cls.tmp, "raw")
        cls.azure = os.path.join(cls.tmp, "azure")
        os.makedirs(cls.raw)
        os.makedirs(cls.azure)
        cls.docx_path = os.path.join(cls.raw, "2567-svc-014.docx")
        cls.azure_path = os.path.join(cls.azure, "2566-nda-002.json")
        make_docx(cls.docx_path)
        make_azure(cls.azure_path)

        cls.out = os.path.join(cls.tmp, "corpus.jtcorpus.gz")
        cls.log = []
        cfg = {"in": cls.raw, "out": cls.out, "work": cls.tmp,
               "azure": cls.azure, "lexicon": os.path.join(HERE, "lexicon")}
        cls.cfg = cfg
        code = pipeline.run(cfg, "all", cls.log.append)
        assert code == 0, "pipeline exited %s\n%s" % (code, "\n".join(cls.log))
        cls.corpus = validate_corpus.load(cls.out)
        cls.blocks = {}
        for path in pipeline.stage_files(cls.tmp, "clean"):
            doc = pipeline.read_stage(path, "blocks")
            cls.blocks[doc["id"]] = doc["blocks"]
        cls.segmented = {}
        for path in pipeline.stage_files(cls.tmp, "segmented"):
            doc = pipeline.read_stage(path, "clauses")
            cls.segmented[doc["id"]] = doc

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.tmp, ignore_errors=True)

    # -- the contract itself --------------------------------------------------

    def test_corpus_validates(self):
        errors = validate_corpus.validate(self.corpus)
        self.assertEqual(errors, [], "\n".join(errors))

    def test_fixture_still_validates(self):
        fixture = os.path.join(HERE, "fixtures", "sample.corpus.json")
        self.assertEqual(validate_corpus.validate(validate_corpus.load(fixture)), [])

    def test_gzip_magic(self):
        with open(self.out, "rb") as handle:
            self.assertEqual(handle.read(2), b"\x1f\x8b")

    def test_b_decodes_to_tokens_that_tile_the_text(self):
        """`b` is the load-bearing encoding: decoded tokens must rejoin to `t`."""
        total = 0
        for index, clause in enumerate(self.corpus["clauses"]):
            tokens = annotate.decode_tokens(clause["t"], clause["b"])
            self.assertEqual("".join(tokens), clause["t"], "clause %d" % index)
            total += len(tokens)
        self.assertEqual(total, self.corpus["stats"]["tokens"])

    def test_doc_ranges_tile_clauses(self):
        cursor = 0
        for doc in self.corpus["docs"]:
            self.assertEqual(doc["c"][0], cursor)
            cursor = doc["c"][1]
        self.assertEqual(cursor, len(self.corpus["clauses"]))

    # -- ingest ---------------------------------------------------------------

    def test_both_ingesters_ran(self):
        methods = sorted(d["src"]["method"] for d in self.corpus["docs"])
        self.assertEqual(methods, ["azure-di:prebuilt-layout", "docx"])

    def test_no_block_is_dropped(self):
        """Every block with text is in a clause, or is a heading that became `h`."""
        for doc_id, blocks in self.blocks.items():
            placed = set()
            headings = set()
            for clause in self.segmented[doc_id]["clauses"]:
                placed.update(clause.get("blocks") or [])
                if clause.get("h"):
                    headings.add(clause["h"].strip())
            for block in blocks:
                if not (block.get("text") or "").strip():
                    continue
                if block["i"] in placed:
                    continue
                self.assertIn(block["text"].strip(), headings,
                              "%s block %d was dropped: %r"
                              % (doc_id, block["i"], block["text"][:60]))

    def test_table_rows_survive_in_document_order(self):
        for doc_id, doc in self.segmented.items():
            rows = [c for c in doc["clauses"] if c.get("srcKind") == "row"]
            self.assertTrue(rows, "no table row in %s" % doc_id)
            self.assertTrue(any(" | " in c["t"] for c in rows))
        docx_clauses = self.segmented["2567-svc-014"]["clauses"]
        row_at = next(i for i, c in enumerate(docx_clauses) if c.get("srcKind") == "row")
        after = " ".join(c["t"] for c in docx_clauses[row_at:])
        self.assertIn("บอกเลิกสัญญา", after,
                      "the table must sit before the termination clause, not at the end")

    def test_azure_roles_are_used_not_re_derived(self):
        azure = self.segmented["2566-nda-002"]["clauses"]
        marginal = [c for c in azure if c.get("marginal")]
        self.assertEqual(len(marginal), 2)          # pageHeader + pageFooter
        self.assertTrue(all(c["z"] == "unplaced" for c in marginal))
        headings = [c for c in azure if c.get("srcKind") == "heading"]
        self.assertTrue(any("การรักษาความลับ" in (c.get("h") or "") for c in headings))

    def test_azure_cell_paragraphs_are_not_emitted_twice(self):
        texts = [c["t"] for c in self.segmented["2566-nda-002"]["clauses"]]
        self.assertEqual(sum(1 for t in texts if "แผนธุรกิจและงบการเงิน" in t), 1)

    def test_azure_confidence_is_carried_through(self):
        doc = next(d for d in self.corpus["docs"] if d["src"]["method"].startswith("azure"))
        self.assertIsNotNone(doc["qc"]["meanConf"])
        self.assertLess(doc["qc"]["meanConf"], 1.0)
        self.assertGreaterEqual(doc["qc"]["lowConfBlocks"], 1)
        flagged = [c for c in self.corpus["clauses"]
                   if c.get("q") and c["d"] == self.corpus["docs"].index(doc)]
        self.assertTrue(flagged, "a low-confidence block must reach the corpus flagged")

    # -- numbering ------------------------------------------------------------

    def test_docx_numbers_come_from_numbering_xml(self):
        clauses = self.segmented["2567-svc-014"]["clauses"]
        labels = {c["nd"]: c for c in clauses if c["nd"]}
        self.assertIn("1.", labels)                 # parties, decimal
        self.assertIn("ก.", labels)                 # recitals, thaiLetters
        self.assertIn("(1)", labels)                # third level, parenthesised
        self.assertEqual(labels["ก."]["n"], "ก")
        self.assertEqual(labels["ก."]["p"], [1])    # ordinals, not text (§4.3)
        self.assertEqual(labels["ข."]["p"], [2])
        deep = [c for c in clauses if len(c["p"]) == 2]
        self.assertTrue(deep)
        self.assertEqual(deep[0]["n"], "1.1")

    def test_docx_numbering_restarts_mid_document(self):
        clauses = self.segmented["2567-svc-014"]["clauses"]
        firsts = [c["nd"] for c in clauses if c["nd"] in ("1.", "ก.")]
        self.assertEqual(firsts[:3], ["1.", "ก.", "1."],
                         "parties 1-2, recitals ก-ข, then the operative clauses restart at 1.")

    def test_azure_numbers_come_from_the_text_regex(self):
        clauses = self.segmented["2566-nda-002"]["clauses"]
        numbered = {c["n"]: c for c in clauses if c["n"]}
        self.assertIn("1.1", numbered)              # ๑.๑ normalized to Arabic
        self.assertEqual(numbered["1.1"]["nd"], "๑.๑")
        self.assertEqual(numbered["1.1"]["p"], [1, 1])
        self.assertFalse(numbered["1.1"]["t"].startswith("๑.๑"),
                         "the literal number is moved into nd, not left in t")

    def test_custom_heading_style_is_recognised(self):
        blocks = self.blocks["2567-svc-014"]
        title = next(b for b in blocks if b.get("style") == "CenterHeading")
        self.assertEqual(title["kind"], "heading")

    # -- clean, segment, annotate --------------------------------------------

    def test_thai_spaces_are_kept(self):
        text = next(c["t"] for c in self.corpus["clauses"] if "บอกเลิกสัญญาได้ทันที" in c["t"])
        self.assertIn(" ", text, "the space is a phrase boundary and carries information")

    def test_zones_are_assigned(self):
        zones = {c["z"] for c in self.corpus["clauses"]}
        for expected in ("preamble", "parties", "body", "signature", "unplaced"):
            self.assertIn(expected, zones)

    def test_signature_clauses_are_indexed_not_deleted(self):
        signature = [c for c in self.corpus["clauses"] if c["z"] == "signature"]
        self.assertTrue(signature)
        self.assertTrue(all(c["b"] for c in signature))

    def test_languages(self):
        langs = {c["l"] for c in self.corpus["clauses"]}
        self.assertIn("th", langs)
        self.assertIn("en", langs)
        english = next(c for c in self.corpus["clauses"] if "Receiving Party" in c["t"])
        self.assertEqual(english["l"], "en")

    def test_clause_kinds(self):
        kinds = {c["k"] for c in self.corpus["clauses"] if c["k"]}
        self.assertIn("termination", kinds)
        self.assertIn("governing_law", kinds)
        for clause in self.corpus["clauses"]:
            if clause["k"]:
                self.assertIn(clause["k"], self.corpus["kinds"])

    def test_duplicates_are_kept_and_counted(self):
        """The same governing-law boilerplate is in both documents."""
        hits = [c for c in self.corpus["clauses"] if GOVERNING_LAW in c["t"]]
        self.assertEqual(len(hits), 2, "duplicates are kept, never removed")
        self.assertTrue(all(c["u"] == 2 for c in hits), "and the count is the signal")

    def test_document_metadata(self):
        docx_doc = next(d for d in self.corpus["docs"] if d["src"]["method"] == "docx")
        self.assertEqual(docx_doc["date"], "2024-03-11")     # ๑๑ มีนาคม ๒๕๖๗, BE → ISO
        self.assertEqual(docx_doc["type"], "service")
        self.assertEqual(docx_doc["lang"], "th")
        self.assertIn("ผู้ว่าจ้าง", [p["role"] for p in docx_doc["parties"]])
        azure_doc = next(d for d in self.corpus["docs"] if d["src"]["method"] != "docx")
        self.assertEqual(azure_doc["date"], "2023-11-02")
        self.assertEqual(azure_doc["type"], "nda")

    def test_frequency_table_excludes_repeating_furniture(self):
        th = dict(self.corpus["freq"]["th"])
        self.assertGreater(th.get("สัญญา", 0), 0)
        self.assertNotIn("จาก", [t for t, _ in self.corpus["freq"]["th"] if t == "จาก"][:0])
        header_word = "ฉบับลงนาม"
        self.assertNotIn(header_word, th, "page headers must not feed the frequency table")

    def test_stop_lists_ship_with_the_corpus(self):
        self.assertIn("และ", self.corpus["stop"]["th"])
        self.assertIn("the", self.corpus["stop"]["en"])

    # -- QC -------------------------------------------------------------------

    def test_manifest(self):
        path = os.path.join(self.tmp, "manifest.csv")
        self.assertTrue(os.path.exists(path))
        with open(path, encoding="utf-8") as handle:
            rows = handle.read().strip().split("\n")
        self.assertEqual(len(rows), 3)              # header + two documents
        self.assertIn("oovRate", rows[0])

    def test_qc_reports_and_never_rejects(self):
        for doc in self.corpus["docs"]:
            for key in ("thaiRatio", "oovRate", "meanConf", "lowConfBlocks",
                        "unplacedBlocks"):
                self.assertIn(key, doc["qc"])
            self.assertLessEqual(doc["qc"]["oovRate"], 1.0)
            self.assertGreaterEqual(doc["qc"]["thaiRatio"], 0.5)

    # -- the CLI, stage by stage ---------------------------------------------

    def test_each_stage_is_separately_invocable(self):
        work = tempfile.mkdtemp(prefix="jtcorpus-stages-")
        try:
            out = os.path.join(work, "corpus.jtcorpus.json")
            base = ["--in", self.raw, "--azure", self.azure, "--work", work, "--out", out]
            for stage in pipeline.STAGES:
                code = pipeline.main(["--stage", stage] + base)
                self.assertEqual(code, 0, "stage %s exited %s" % (stage, code))
            corpus = validate_corpus.load(out)       # uncompressed also loads (§4)
            self.assertEqual(validate_corpus.validate(corpus), [])
            self.assertEqual(corpus["stats"]["clauses"], self.corpus["stats"]["clauses"])
        finally:
            shutil.rmtree(work, ignore_errors=True)

    # -- the non-negotiable ---------------------------------------------------

    def test_no_network_anywhere_in_the_pipeline(self):
        """DESIGN.md §1.1 — the pipeline never calls Azure or anything else."""
        forbidden = ("import requests", "import urllib", "from urllib",
                     "import socket", "import http", "from http",
                     "httpx", "aiohttp", "urlopen", "DocumentIntelligence")
        for name in sorted(os.listdir(HERE)):
            if not name.endswith(".py"):
                continue
            with open(os.path.join(HERE, name), encoding="utf-8") as handle:
                source = handle.read()
            for needle in forbidden:
                self.assertNotIn(needle, source, "%s mentions %s" % (name, needle))


if __name__ == "__main__":
    unittest.main(verbosity=2)
